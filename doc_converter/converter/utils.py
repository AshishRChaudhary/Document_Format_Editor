import os
from pdf2docx import Converter
from docx import Document
from reportlab.pdfgen import canvas
from openpyxl import Workbook
import csv
from PIL import Image

def convert_file(file_path, target_format):
    base, ext = os.path.splitext(file_path)
    converted_path = f"{base}.{target_format}"
    
    if ext == ".pdf" and target_format == "docx":
        cv = Converter(file_path)
        cv.convert(converted_path)
        cv.close()
    elif ext == ".docx" and target_format == "pdf":
        doc = Document(file_path)
        pdf = canvas.Canvas(converted_path)
        for para in doc.paragraphs:
            pdf.drawString(100, 800, para.text)
        pdf.save()
    elif ext == ".txt" and target_format == "pdf":
        with open(file_path, "r") as txt_file:
            pdf = canvas.Canvas(converted_path)
            for i, line in enumerate(txt_file.readlines()):
                pdf.drawString(100, 800 - (i * 20), line.strip())
            pdf.save()
    elif ext == ".png" and target_format == "pdf":
        image = Image.open(file_path)
        image.convert("RGB").save(converted_path)
    elif ext == ".xlsx" and target_format == "pdf":
        pdf = canvas.Canvas(converted_path)
        pdf.drawString(100, 800, "Excel file converted to PDF (Dummy content)")
        pdf.save()
    elif ext == ".xlsx" and target_format == "docx":
        doc = Document()
        doc.add_paragraph("Excel file converted to DOCX (Dummy content)")
        doc.save(converted_path)
    elif ext == ".csv" and target_format == "pdf":
        with open(file_path, newline='') as csvfile:
            reader = csv.reader(csvfile)
            pdf = canvas.Canvas(converted_path)
            y = 800
            for row in reader:
                pdf.drawString(100, y, ", ".join(row))
                y -= 20
            pdf.save()
    elif ext == ".csv" and target_format == "docx":
        doc = Document()
        with open(file_path, newline='') as csvfile:
            reader = csv.reader(csvfile)
            for row in reader:
                doc.add_parag
